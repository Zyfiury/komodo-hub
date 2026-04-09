"""
Add appendix cross-references, in-text citations, missing appendices H–J,
and reference-list entries. Reads ENHANCED (or SUBMIT_READY_edited) and writes
Omar_Zakhama_14498572_SUBMIT_READY_CITATIONS.docx

Run: python report_docs/finalize_report_citations.py
"""
from __future__ import annotations

import glob
import os
import shutil
import time

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def pick_source_docx() -> str:
    base = os.path.dirname(__file__)
    cands = glob.glob(os.path.join(base, "Omar_Zakhama_14498572_SUBMIT_READY*.docx"))
    skip = ("BACKUP", "CITATIONS", "SUBMISSION")
    cands = [c for c in cands if os.path.isfile(c) and not any(s in c for s in skip)]
    if not cands:
        raise FileNotFoundError("No suitable Omar_Zakhama*SUBMIT_READY*.docx (use ENHANCED or edited).")
    return max(cands, key=os.path.getmtime)


def main() -> None:
    src = pick_source_docx()
    out = os.path.join(
        os.path.dirname(__file__),
        "Omar_Zakhama_14498572_SUBMIT_READY_CITATIONS.docx",
    )
    copied = False
    for attempt in range(6):
        try:
            shutil.copy2(src, out)
            copied = True
            break
        except OSError:
            if attempt < 5:
                time.sleep(0.45)
            else:
                alt = out.replace(".docx", "_STAGING.docx")
                shutil.copy2(src, alt)
                print("Could not overwrite CITATIONS (file may be open in Word). Wrote:", alt)
                out = alt
                copied = True
    if not copied:
        raise RuntimeError("Failed to copy source to CITATIONS output.")
    doc = Document(out)

    for p in doc.paragraphs:
        if "\ufffd" in p.text:
            p.text = p.text.replace("\ufffd", "'")

    # --- Link body to appendices (Web Developer week + interface summary) ---
    for p in doc.paragraphs:
        if "Selected website prototype screens are provided in Appendices A to D" in p.text:
            if "Appendices E" not in p.text:
                p.text = (
                    p.text.rstrip()
                    + " Supplementary marker-facing materials (repository URL, local execution, seeded role logins, "
                    "and contact for access issues) are provided in Appendices E–H, following the coursework access "
                    "requirements (Talekar, 2025c; Talekar, 2025d). AI use and originality declarations appear in "
                    "Appendices I–J."
                )
        if p.text.startswith("Figures in Appendices") and "repository README" in p.text:
                p.text = (
                    "Figures in Appendices A to D provide screenshot evidence aligned with the prototype narrative above "
                    "(Hapsara, 2023). Step-by-step execution for assessors is in Appendix F; credentials in Appendix G "
                    "(Talekar, 2025c)."
                )

    # --- Cloud / architecture sentences ---
    for p in doc.paragraphs:
        if p.text.startswith("The deployment view uses managed cloud services"):
            if "Appendix E" not in p.text:
                p.text += " Further repository and configuration notes appear in Appendix E."
        if p.text.startswith("The logical architecture isolates presentation"):
            if "Appendix E" not in p.text:
                p.text += " Service boundaries are reflected in the repository layout described in Appendix E."

    # --- Testing block: cite brief + ISO where we mention verification ---
    for p in doc.paragraphs:
        if p.text.startswith("Automated tests: eslint"):
            if "Talekar, 2025d" not in p.text:
                p.text += " Evidence aligns with the quality and documentation expectations in the assignment brief (Talekar, 2025d)."

    # --- Conclusion cross-reference ---
    for p in doc.paragraphs:
        if p.text.startswith("This report outlines a full eight-week Scrum cycle"):
            if "Evidence packaging for assessment" not in p.text:
                p.text += (
                    " Evidence packaging for assessment follows Appendices A to K: user-facing screenshots (A to D), "
                    "repository and runbook (E to F), test accounts (G), support email (H), AI use (I), originality (J), "
                    "and submission checklist (K) (Talekar, 2025c)."
                )

    # --- Appendix figure captions: tie to case study ---
    for p in doc.paragraphs:
        if p.text.startswith("Figure A1.") or p.text.startswith("Figure B1.") or p.text.startswith("Figure C1."):
            if "(Hapsara, 2023)" not in p.text:
                p.text = p.text.rstrip() + " (Hapsara, 2023)."
        if p.text.startswith("Figure D1.") and "(Hapsara, 2023)" not in p.text:
            p.text = p.text.rstrip() + " (Hapsara, 2023)."

    # --- APPENDICES intro ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "APPENDICES":
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
                j += 1
            nxt = doc.paragraphs[j] if j < len(doc.paragraphs) else None
            intro = (
                "The appendices below are referenced from the sprint sections and conclusion. Appendices A to D illustrate the working "
                "prototype; E to G satisfy repository, execution, and credential requirements; H to J cover support, AI use, "
                "and academic integrity; Appendix K is a submission checklist (Talekar, 2025c)."
            )
            if nxt and nxt.text.strip().startswith("Appendix A"):
                if not any("The appendices below are referenced" in x.text for x in doc.paragraphs):
                    insert_paragraph_after(p, intro)
            break

    # --- REFERENCES: web source for stack (optional; remove if your tutor disallows non-scholarly URLs)
    ref_text = "\n".join(p.text for p in doc.paragraphs)
    insert_after_iso = None
    for p in doc.paragraphs:
        if p.text.startswith("International Organisation for Standardization"):
            insert_after_iso = p
            break
    if insert_after_iso and "nextjs.org/docs" not in ref_text:
        insert_paragraph_after(
            insert_after_iso,
            "Vercel (2026) Next.js Documentation. Available at: https://nextjs.org/docs (Accessed: 9 April 2026).",
        )

    # --- Appendices H, I, J if missing ---
    full = "\n".join(p.text for p in doc.paragraphs)
    if "Appendix H. Support contact" not in full:
        anchor = None
        for p in doc.paragraphs:
            if p.text.strip().startswith("Public routes for anonymous checks"):
                anchor = p
                break
        if anchor is None:
            anchor = doc.paragraphs[-1]
        for line in [
            "",
            "Appendix H. Support contact for access issues",
            "Primary student contact (replace with your active university email): [YOUR.COVENTRY.EMAIL@uni.coventry.ac.uk]",
            "Use this inbox for marker questions about registration, codes, or environment errors during evaluation.",
            "",
            "Appendix I. Summary of AI use",
            "Generative AI assisted with drafting clarity, checklist alignment against the template, appendix structuring, "
            "and citation cross-referencing. Technical descriptions were verified against the implemented repository and "
            "module PDFs. I remain accountable for the accuracy and integrity of this submission.",
            "",
            "Appendix J. Declaration of originality",
            "I confirm this report and referenced prototype are my own work except where sources are cited. I understand the "
            "Faculty regulations on academic integrity and the consequences of academic misconduct.",
            "",
        ]:
            anchor = insert_paragraph_after(anchor, line)

    doc.save(out)
    print("Wrote:", out)
    print("Source used:", src)


if __name__ == "__main__":
    main()
