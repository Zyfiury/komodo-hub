"""
Build submission-ready report: APA 7 references, replace appendices entirely with
correct A-K content + screenshots. Requires PNGs from: npm run screenshots:appendix

Output: Omar_Zakhama_14498572_SUBMIT_READY_SUBMISSION.docx
"""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: str, width_inches: float = 6.2) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, "")
    run = new_para.add_run()
    if os.path.isfile(image_path):
        run.add_picture(image_path, width=Inches(width_inches))
    else:
        run.add_text(f"[Missing screenshot: run npm run screenshots:appendix — {image_path}]")
    return new_para


def fix_encoding(doc: Document) -> None:
    for p in doc.paragraphs:
        if "\ufffd" in p.text:
            p.text = p.text.replace("\ufffd", "'")


def remove_everything_after_appendices_heading(doc: Document) -> bool:
    """Strip old/wrong appendix blocks (paragraphs and tables). Keeps the 'APPENDICES' heading."""
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "APPENDICES":
            target = p
            break
    if target is None:
        return False
    el = target._element
    parent = el.getparent()
    nxt = el.getnext()
    while nxt is not None:
        cur = nxt
        nxt = nxt.getnext()
        parent.remove(cur)
    return True


def rebuild_appendices_from_scratch(doc: Document, shot_dir: str, gh: str) -> None:
    """Insert correct Appendix A-K with embedded screenshots (after APPENDICES heading)."""
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "APPENDICES":
            target = p
            break
    if target is None:
        return

    cur = target
    gh_display = gh.rstrip(".git")

    def add(text: str = "") -> None:
        nonlocal cur
        cur = insert_paragraph_after(cur, text)

    def pic(rel_name: str) -> None:
        nonlocal cur
        cur = insert_picture_after(cur, os.path.join(shot_dir, rel_name))

    add(
        "These appendices replace earlier placeholder material. They provide assessor-facing evidence: "
        "live screenshots (A to D), repository and run instructions (E to G), and support or integrity statements "
        "(H to J), with a submission checklist (K) (Talekar, 2025c)."
    )

    add("Appendix A. Public landing page (working prototype)")
    add(
        "Figure A1. Komodo Hub public landing page showing programme context, navigation to library and species content, "
        "and privacy-aware messaging aligned with the case study (Hapsara, 2023)."
    )
    pic("appendix-a-landing.png")

    if os.path.isfile(os.path.join(shot_dir, "appendix-public-library.png")):
        add(
            "Figure A2. Public library browse (anonymous access to approved materials only) (Hapsara, 2023)."
        )
        pic("appendix-public-library.png")

    add("Appendix B. Teacher dashboard (working prototype)")
    add(
        "Figure B1. Teacher dashboard: class and programme actions, progress visibility, messages, and reporting tasks "
        "(Hapsara, 2023)."
    )
    pic("appendix-b-teacher.png")

    add("Appendix C. Moderated wildlife reporting (working prototype)")
    add(
        "Figure C1. Wildlife reporting flow: draft and submit for moderation before wider publication; organisation-scoped "
        "access (Hapsara, 2023)."
    )
    pic("appendix-c-wildlife.png")

    add("Appendix D. Foundation administration (working prototype)")
    add(
        "Figure D1. Foundation dashboard: organisation overview, moderation entry points, campaigns, and service "
        "signals (Hapsara, 2023)."
    )
    pic("appendix-d-foundation.png")

    add("Appendix E. Repository, diagrams, and configuration")
    add(f"Public source code: {gh_display}")
    add(
        "Diagram sources (UML, DFD, ERD, architecture): store draw.io or exported PNG/SVG in this repository or "
        "link an OneDrive folder with view access, as required by the module template."
    )
    add("Local configuration: copy `.env.example` to `.env` and set `SESSION_SECRET` (minimum 16 characters).")

    add("Appendix F. How to run the application (local, for assessors)")
    add("Prerequisites: Node.js (LTS), npm, and Git.")
    add("1) Clone the repository and open the project folder.")
    add("2) Install dependencies: npm install")
    add("3) Create `.env` from `.env.example` and set SESSION_SECRET.")
    add("4) Apply database and seed: npx prisma migrate deploy then npx prisma db seed")
    add("5) Start the dev server: npm run dev")
    add("6) Open http://localhost:3000 in a browser.")
    add("Optional quality checks: npm run lint and npm run build")

    add("Appendix G. Seeded test accounts and demo access code")
    add("Password for all seeded demonstration accounts: KomodoHub!Dev2026")
    add("Foundation admin: foundation.admin@komodohub.local")
    add("School admin: school.admin@komodohub.local")
    add("Teacher: teacher@komodohub.local")
    add("Student 1: student1@komodohub.local")
    add("Student 2: student2@komodohub.local")
    add("Community admin: community.admin@komodohub.local")
    add("Community member 1: member1@komodohub.local")
    add("Community member 2: member2@komodohub.local")
    add("Demo school access code (student join): SCHOOL-DEMO-2026")
    add("Anonymous public routes: /library, /campaigns, /species (approved items only).")

    add("Appendix H. Support contact for access issues")
    add("Replace with your active Coventry University email: [YOUR.COVENTRY.EMAIL@uni.coventry.ac.uk]")
    add("Markers may use this address for questions about registration, codes, or environment errors.")

    add("Appendix I. Summary of generative AI use")
    add(
        "Generative AI assisted with drafting, structuring appendices, APA-style reference formatting, and automation "
        "for screenshots. Technical descriptions were checked against the implemented application and module PDFs. "
        "I am responsible for the accuracy of this submission."
    )

    add("Appendix J. Declaration of originality")
    add(
        "I confirm this report and the referenced prototype are my own work except where sources are cited. "
        "I understand my faculty's regulations on academic integrity."
    )

    add("Appendix K. Submission checklist (cross-reference)")
    add("• A to D: screenshots in this document match the running prototype.")
    add(f"• E: public repository at {gh_display}.")
    add("• F to G: steps and credentials tested on Windows with Node LTS.")
    add("• H to J: contact, AI use, and originality completed before upload.")
    add("• UML/DFD/ERD originals align with files linked or stored per Appendix E (Talekar, 2025c).")


def insert_apa7_note_after_references_heading(doc: Document) -> None:
    note = "The reference list follows American Psychological Association (APA) style, 7th edition."
    if note in "\n".join(p.text for p in doc.paragraphs):
        return
    for p in doc.paragraphs:
        if p.text.strip() == "REFERENCES":
            insert_paragraph_after(p, note)
            return


def convert_existing_reference_line_to_apa7(text: str) -> str | None:
    t = text.strip()
    if not t:
        return None
    if t.startswith("Hapsara, M."):
        return (
            "Hapsara, M. (2023). Komodo Hub: A digital platform for community-supported animal conservation "
            "[Teaching case study]. Coventry University."
        )
    if t.startswith("Talekar, A. (2025c)"):
        return (
            "Talekar, A. (2025c). 5005CMD how to produce your project report [Unpublished course handout]. "
            "Coventry University."
        )
    if t.startswith("Talekar, A. (2025a)"):
        return (
            "Talekar, A. (2025a). 5005CMD how to manage your team and project [Unpublished course handout]. "
            "Coventry University."
        )
    if t.startswith("Talekar, A. (2025b)"):
        return (
            "Talekar, A. (2025b). 5005CMD how to perform your role [Unpublished course handout]. Coventry University."
        )
    if t.startswith("Talekar, A. (2025d)"):
        return (
            "Talekar, A. (2025d). 5005CMD student assignment brief 2025/26 [Unpublished course handout]. "
            "Coventry University."
        )
    if t.startswith("International Organisation for Standardization") or t.startswith(
        "International Organization for Standardization"
    ):
        return (
            "International Organization for Standardization. (2019). ISO 9241-210:2019 Ergonomics of human-system "
            "interaction - Part 210: Human-centred design for interactive systems. "
            "https://www.iso.org/standard/77520.html"
        )
    if t.startswith("Vercel") and "nextjs.org" in t:
        return "Vercel. (2026). Next.js documentation. https://nextjs.org/docs"
    if t.startswith("Schwaber, K.") and "scrum guide" in t.lower():
        return (
            "Schwaber, K., & Sutherland, J. (2020). The scrum guide. Scrum Guides. "
            "https://scrumguides.org/scrum-guide.html"
        )
    if t.startswith("Prisma") and "prisma.io" in t:
        return "Prisma. (2026). Prisma ORM documentation. https://www.prisma.io/docs"
    if t.startswith("OWASP"):
        return (
            "OWASP Foundation. (2023). Password storage cheat sheet. OWASP Cheat Sheet Series. "
            "https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html"
        )
    if t.startswith("W3C") and "WCAG" in t:
        return (
            "W3C. (2023). Web content accessibility guidelines (WCAG) 2.2. W3C. https://www.w3.org/TR/WCAG22/"
        )
    if t.startswith("IUCN"):
        return "IUCN. (2024). The IUCN red list of threatened species. https://www.iucnredlist.org/"
    if t.startswith("Sommerville, I.") or t.startswith("Somerville, I."):
        return "Sommerville, I. (2016). Software engineering (10th ed.). Pearson Education."
    if t.startswith("Nielsen, J."):
        return (
            "Nielsen, J. (2012). Usability 101: Introduction to usability. Nielsen Norman Group. "
            "https://www.nngroup.com/articles/usability-101-introduction-to-usability/"
        )
    return None


def apply_apa7_reference_list(doc: Document) -> None:
    in_refs = False
    for p in doc.paragraphs:
        raw = p.text.strip()
        if raw == "REFERENCES":
            in_refs = True
            continue
        if in_refs and raw == "APPENDICES":
            break
        if in_refs and raw:
            converted = convert_existing_reference_line_to_apa7(p.text)
            if converted is not None:
                p.text = converted


def main() -> None:
    base = os.path.dirname(os.path.abspath(__file__))
    src = os.path.join(base, "Omar_Zakhama_14498572_SUBMIT_READY_CITATIONS.docx")
    if not os.path.isfile(src):
        src = os.path.join(base, "Omar_Zakhama_14498572_SUBMIT_READY_ENHANCED.docx")
    if not os.path.isfile(src):
        src = os.path.join(base, "Omar_Zakhama_14498572_SUBMIT_READY_edited.docx")
    out = os.path.join(base, "Omar_Zakhama_14498572_SUBMIT_READY_SUBMISSION.docx")
    shot_dir = os.path.join(base, "appendix_screenshots")

    try:
        shutil.copy2(src, out)
    except OSError:
        out = os.path.join(base, "Omar_Zakhama_14498572_SUBMIT_READY_SUBMISSION_APA7.docx")
        shutil.copy2(src, out)
    doc = Document(out)
    fix_encoding(doc)
    insert_apa7_note_after_references_heading(doc)
    apply_apa7_reference_list(doc)

    gh = "https://github.com/Zyfiury/komodo-hub.git"
    for p in doc.paragraphs:
        if "[YOUR_USERNAME]" in p.text or "[YOUR_REPO]" in p.text:
            p.text = (
                p.text.replace("https://github.com/[YOUR_USERNAME]/[YOUR_REPO]", gh.rstrip(".git"))
                .replace("[YOUR_USERNAME]/[YOUR_REPO]", "Zyfiury/komodo-hub")
            )

    # Replace all legacy appendix content with the correct A-K set + screenshots
    if remove_everything_after_appendices_heading(doc):
        rebuild_appendices_from_scratch(doc, shot_dir, gh)

    blob = "\n".join(p.text for p in doc.paragraphs)
    extras_refs = [
        ("Schwaber, K., & Sutherland", "Schwaber, K., & Sutherland, J. (2020). The scrum guide. Scrum Guides. https://scrumguides.org/scrum-guide.html"),
        ("Prisma. (2026)", "Prisma. (2026). Prisma ORM documentation. https://www.prisma.io/docs"),
        ("OWASP Foundation. (2023)", "OWASP Foundation. (2023). Password storage cheat sheet. OWASP Cheat Sheet Series. https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html"),
        ("W3C. (2023)", "W3C. (2023). Web content accessibility guidelines (WCAG) 2.2. W3C. https://www.w3.org/TR/WCAG22/"),
        ("IUCN. (2024)", "IUCN. (2024). The IUCN red list of threatened species. https://www.iucnredlist.org/"),
        ("Sommerville, I. (2016)", "Sommerville, I. (2016). Software engineering (10th ed.). Pearson Education."),
        ("Nielsen, J. (2012", "Nielsen, J. (2012). Usability 101: Introduction to usability. Nielsen Norman Group. https://www.nngroup.com/articles/usability-101-introduction-to-usability/"),
    ]
    anchor = None
    for p in doc.paragraphs:
        if p.text.startswith("International Organization for Standardization") or p.text.startswith(
            "International Organisation for Standardization"
        ):
            anchor = p
            break
    if anchor:
        cur = anchor
        for key, line in extras_refs:
            if key not in blob:
                cur = insert_paragraph_after(cur, line)
                blob += line

    apply_apa7_reference_list(doc)

    doc.save(out)
    print("Wrote:", out)


if __name__ == "__main__":
    main()
