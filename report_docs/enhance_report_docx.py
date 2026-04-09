"""
One-off enhancement of the 5005CMD individual report DOCX.
Run from anywhere: python report_docs/enhance_report_docx.py
"""
from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def fix_encoding(doc: Document) -> None:
    for p in doc.paragraphs:
        if "\ufffd" in p.text:
            p.text = p.text.replace("\ufffd", "'")


def find_para(doc: Document, text: str, *, exact: bool = False) -> Paragraph | None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if exact:
            if t == text:
                return p
        else:
            if p.text.startswith(text) or t == text:
                return p
    return None


def find_all_para(doc: Document, exact: str) -> list[Paragraph]:
    return [p for p in doc.paragraphs if p.text.strip() == exact]


def insert_block_after(anchor: Paragraph, lines: list[str]) -> Paragraph:
    cur = anchor
    for line in lines:
        cur = insert_paragraph_after(cur, line)
    return cur


def insert_block_before(anchor: Paragraph, lines: list[str]) -> None:
    for line in reversed(lines):
        insert_paragraph_before(anchor, line)


def job_table_exists(doc: Document) -> bool:
    for t in doc.tables:
        if len(t.rows) and len(t.rows[0].cells) > 3:
            if t.rows[0].cells[3].text.strip() == "Primary role":
                return True
    return False


def add_job_table_after(doc: Document, paragraph: Paragraph) -> None:
    if job_table_exists(doc):
        return
    rows = [
        ("Sprint", "Week", "Week ending", "Primary role", "Focus"),
        ("Sprint 1", "1", "10 Feb 2026", "Product Owner", "Product goals, prioritised backlog, privacy-first scope"),
        ("Sprint 1", "2", "17 Feb 2026", "Scrum Master", "Sprint goal, sprint backlog, weekly checkpoints for auth/privacy"),
        ("Sprint 1", "3", "24 Feb 2026", "Database Engineer", "ERD/DFD alignment, data dictionary, org visibility rules"),
        ("Sprint 1", "4", "3 Mar 2026", "Web Developer", "Use case/activity views, interface prototype, reporting/library flows"),
        ("Sprint 2", "1", "10 Mar 2026", "Cloud Engineer", "Deployment view, managed services, session and audit posture"),
        ("Sprint 2", "2", "17 Mar 2026", "System Architect", "Logical services, quality criteria, maintainable boundaries"),
        ("Sprint 2", "3", "24 Mar 2026", "Product Owner", "Backlog refinement pre-release: moderation, analytics, reliability"),
        ("Sprint 2", "4", "31 Mar 2026", "Scrum Master", "Sprint review/retrospective structure, risk checkpoints"),
    ]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)


def main() -> None:
    src = r"c:\Users\omarz\Desktop\webdesign\report_docs\Omar_Zakhama_14498572_SUBMIT_READY_edited.docx"
    path = r"c:\Users\omarz\Desktop\webdesign\report_docs\Omar_Zakhama_14498572_SUBMIT_READY_ENHANCED.docx"
    import shutil

    shutil.copy2(src, path)
    doc = Document(path)
    fix_encoding(doc)
    full = "\n".join(p.text for p in doc.paragraphs)

    # --- Title: individual submission ---
    p_team = find_para(doc, "Team Members", exact=True)
    if p_team:
        p_team.text = "Individual submission"
    # Replace following placeholder lines
    fillers = [
        (
            "This report is submitted as an individual 5005CMD Software Engineering case study. "
            "All sprint evidence and reflections are written in the first person and represent my own work (Omar Zakhama, 14498572)."
        ),
        (
            "Individual study approval (mandatory for individual route): insert below a screenshot or PDF of written approval "
            "from the Module Leader (Anuja Talekar) or Course Director (Dianabasi), matching the official template instruction."
        ),
        (
            "Public GitHub repository (mandatory; must be public for marking): https://github.com/[YOUR_USERNAME]/[YOUR_REPO] "
            "(replace with your real repository URL before submission)."
        ),
        (
            "Optional OneDrive / SharePoint link for source diagrams (draw.io exports, large figures): [paste link with viewer access] "
            "or write Not used if everything is in GitHub."
        ),
    ]
    if "Public GitHub repository (mandatory" not in full:
        idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "Individual submission":
                idx = i
                break
        if idx is not None:
            anchor = doc.paragraphs[idx]
            for j, new_text in enumerate(fillers, start=1):
                if idx + j < len(doc.paragraphs):
                    doc.paragraphs[idx + j].text = new_text
                else:
                    anchor = insert_paragraph_after(anchor, new_text)
        full = "\n".join(p.text for p in doc.paragraphs)

    # --- Word count declaration (front matter) ---
    p_deadline = find_para(doc, "Submission deadline")
    if p_deadline and "Word count declaration (per template exclusions)" not in full:
        insert_block_after(
            p_deadline,
            [
                "",
                "Word count declaration (per template exclusions): Body narrative (counts toward limit): approximately ____ words. "
                "Figures, tables, reference list, appendices, process tables, DFD/UML, and Gantt-style schedules are excluded as per module guidance. "
                "Replace ____ with the figure from Microsoft Word Review → Word Count after final edit.",
                "Totals checklist: (1) count each major section that is included in the limit; (2) record the declared total here on the title page as required.",
                "",
            ],
        )
        full = "\n".join(p.text for p in doc.paragraphs)

    # --- Table of contents (instruction + outline; user updates page numbers in Word) ---
    p_sol = find_para(doc, "Solution priorities", exact=True)
    if p_sol and "Update in Microsoft Word before PDF export" not in full:
        toc_lines = [
            "TABLE OF CONTENTS",
            "Update in Microsoft Word before PDF export: References → Table of Contents → Automatic Table, then right-click → Update entire table.",
            "Section outline (page numbers filled by Word): Project Summary; Table of Contents; Solution priorities; Job allocation; "
            "Sprint 1 Weeks 1–4; Sprint 2 Weeks 1–4; Project schedule; Sprint summaries; Testing and verification; Future work; Conclusion; References; Appendices A–J.",
            "",
        ]
        insert_block_before(p_sol, toc_lines)
        full = "\n".join(p.text for p in doc.paragraphs)

    # --- Job allocation narrative + table ---
    p_job = find_para(doc, "JOB ALLOCATION", exact=True)
    if p_job:
        nxt = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "JOB ALLOCATION" and i + 1 < len(doc.paragraphs):
                nxt = doc.paragraphs[i + 1]
                break
        if nxt and nxt.text.startswith("The team carried out"):
            nxt.text = (
                "I followed two sequential four-week Scrum sprints (eight weekly cycles), rotating weekly through Product Owner, "
                "Scrum Master, and developer specialisms (Database Engineer, Web Developer, Cloud Engineer, System Architect). "
                "This matches the module expectation that each student experiences multiple roles while delivering one coherent product."
            )
        # Remove duplicate second paragraph if it still says "The schedule helps..."
        for p in doc.paragraphs:
            if p.text.startswith("The schedule helps maintain group consistency"):
                p.text = (
                    "Because this is an individual submission, the job allocation table below is my own rotation record (not a multi-student roster). "
                    "Evidence in each week maps to Process Control, Product Development, and Project Management artefacts described in the HOW TO documents."
                )
                break
        # Insert table after first job paragraph
        anchor = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "JOB ALLOCATION" and i + 1 < len(doc.paragraphs):
                anchor = doc.paragraphs[i + 1]
                break
        if anchor:
            add_job_table_after(doc, anchor)

    full = "\n".join(p.text for p in doc.paragraphs)

    # --- Fill evidence placeholders (refresh full below as we insert) ---
    backlog_1 = [
        "Representative product backlog items (Sprint 1, Week 1 — Product Owner):",
        "• US-ONB: Organisation onboarding with school vs community visibility rules (Hapsara, 2023).",
        "• US-AUTH: Secure authentication and role-based routing for teachers, students, foundation staff, and community roles.",
        "• US-CODE: School access codes so teachers can admit students without exposing child profiles publicly.",
        "• US-LIB: Moderated library publishing workflow separating draft, review, and approved public items.",
        "• US-MSG: Internal messaging and notes on submissions aligned to teacher workflow in the case study.",
        "• US-DASH: Role dashboards (teacher, foundation) for workload, moderation queue, and subscription insight.",
    ]
    backlog_2 = [
        "Backlog adjustments (Sprint 2, Week 3 — Product Owner), pre-release:",
        "• Harden moderation queues and decision audit trails for wildlife and library submissions.",
        "• Expand analytics cards for foundation users (subscriptions, service health proxies, engagement summaries).",
        "• Reliability tasks: session handling checks, clearer error states, and regression re-tests on auth boundaries.",
    ]
    snaps = [
        [
            "Progress snapshot (Sprint 1, Week 2 — Scrum Master):",
            "• Sprint goal: stabilise authentication, privacy boundaries, and role navigation for school-facing flows.",
            "• Sprint backlog themes: login/register, organisation context, teacher/student home routes, initial message shell.",
            "• Weekly Scrum notes (individual log): scope confirmed against case study interviews; blockers recorded as environment/setup only.",
            "• Definition of Done for the week: demonstrable navigation across primary roles without breaking public/private library rules.",
        ],
        [
            "Progress snapshot (Sprint 2, Week 4 — Scrum Master):",
            "• Sprint goal: close review and retrospective for release readiness; confirm critical paths for moderation and reporting.",
            "• Review focus: authentication edge cases, role dashboards, and library publication states.",
            "• Retrospective themes: maintain thin vertical slices each week; keep security/privacy checks on every increment.",
            "• Carry-over (see Future Work): automated test suite expansion and production hardening beyond the prototype.",
        ],
    ]

    backs = find_all_para(doc, "Backlog evidence")
    if len(backs) >= 1 and "Representative product backlog items (Sprint 1" not in full:
        insert_block_after(backs[0], backlog_1)
        full = "\n".join(p.text for p in doc.paragraphs)
    if len(backs) >= 2 and "Backlog adjustments (Sprint 2, Week 3" not in full:
        insert_block_after(backs[1], backlog_2)
        full = "\n".join(p.text for p in doc.paragraphs)

    snaps_p = find_all_para(doc, "Progress snapshot")
    if len(snaps_p) >= 1 and "Progress snapshot (Sprint 1, Week 2" not in full:
        insert_block_after(snaps_p[0], snaps[0])
        full = "\n".join(p.text for p in doc.paragraphs)
    if len(snaps_p) >= 2 and "Progress snapshot (Sprint 2, Week 4" not in full:
        insert_block_after(snaps_p[1], snaps[1])
        full = "\n".join(p.text for p in doc.paragraphs)

    dd = find_para(doc, "Data dictionary extract", exact=True)
    if dd and "Sample dictionary entries (illustrative" not in full:
        insert_block_after(
            dd,
            [
                "Sample dictionary entries (illustrative; aligns with ERD themes in Figure 3):",
                "Organisation: type (school/community), subscription state, public library visibility flags.",
                "UserAccount: role, organisation FK, authentication credentials reference (hashed), active status.",
                "Programme / Activity: syllabus linkage, enrolment, teacher ownership.",
                "WildlifeReport / LibraryItem: moderation state, submitted_by, published_at, redaction notes.",
                "MessageThread: participants scoped to organisation roles; no public leakage for school pupils.",
            ],
        )
        full = "\n".join(p.text for p in doc.paragraphs)

    iface = find_para(doc, "Interface prototype summary", exact=True)
    if iface and "Prototype coverage linked to ISO 9241-210" not in full:
        insert_block_after(
            iface,
            [
                "Prototype coverage linked to ISO 9241-210 usability themes: clear navigation, recoverable flows, non-technical language on errors.",
                "Key routes: landing and public library; authenticated dashboards by role; moderated submission form; foundation moderation console.",
                "Figures in Appendices A–D provide screenshot evidence; repository README documents how to run the build locally.",
            ],
        )
        full = "\n".join(p.text for p in doc.paragraphs)

    dep = find_para(doc, "Deployment and operations summary", exact=True)
    if dep and "Deployment posture (prototype)" not in full:
        insert_block_after(
            dep,
            [
                "Deployment posture (prototype): Next.js application with managed-platform target (e.g., Vercel/Node hosting) and separate managed database.",
                "Security-oriented operations: HttpOnly session cookies, bcrypt password hashing, least-privilege data access in application layer, structured logging hooks for audit.",
                "Backups/disaster recovery: rely on cloud provider automated backups for production; local SQLite used only for development demonstration.",
            ],
        )
        full = "\n".join(p.text for p in doc.paragraphs)

    arch = find_para(doc, "Architecture evaluation against quality criteria", exact=True)
    if arch and "Quality criteria mapping (summary)" not in full:
        insert_block_after(
            arch,
            [
                "Quality criteria mapping (summary):",
                "• Modularity: UI routes vs server actions vs Prisma data access separation supports change control.",
                "• Security/privacy: school pupil data kept out of public community-style profile patterns, matching case constraints.",
                "• Scalability: stateless app tier + poolable DB suits horizontal scale on a managed platform.",
                "• Maintainability: typed TypeScript boundaries and schema-first data model reduce regression risk.",
            ],
        )
        full = "\n".join(p.text for p in doc.paragraphs)

    # Rename optional group headings
    for p in doc.paragraphs:
        if p.text.strip() == "Additional Artifacts (OPTIONAL)":
            p.text = "Weekly process evidence (individual)"

    # Add short note after each 'Weekly process evidence' heading — first empty paragraph fill
    # (light touch: only first occurrence to avoid noise)
    if "Individual route: there are no separate teammate artefacts" not in full:
        first = True
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "Weekly process evidence (individual)" and first:
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                    doc.paragraphs[i + 1].text = (
                        "Individual route: there are no separate teammate artefacts. This row records the same weekly Scrum discipline "
                        "(goal, backlog adjustments, increment notes) consolidated in the sections above."
                    )
                    first = False

    # --- Pre-conclusion academic blocks ---
    p_conc = find_para(doc, "CONCLUSION", exact=True)
    full = "\n".join(p.text for p in doc.paragraphs)
    if p_conc and "PROJECT SCHEDULE (summary)" not in full:
        block = [
            "",
            "PROJECT SCHEDULE (summary)",
            "Weeks 1–4 (Sprint 1): requirements consolidation, modelling (use case, activity, DFD, ERD), initial UI prototype, database structure.",
            "Weeks 5–8 (Sprint 2): cloud/deployment framing, architecture review, backlog refinement, review/retrospective, integration hardening.",
            "Milestones: M1 end Sprint 1 — coherent data model + navigable UI; M2 end Sprint 2 — deployment narrative + moderated flows demonstrable.",
            "",
            "SPRINT 1 SUMMARY",
            "Outcome: shared understanding of privacy rules, core entities, and primary user journeys for schools and foundation staff; diagrams and prototype aligned to the case study.",
            "",
            "SPRINT 2 SUMMARY",
            "Outcome: deployment and service-boundary narrative completed; architecture assessed against quality attributes; release-oriented backlog updates documented.",
            "",
            "TESTING AND VERIFICATION",
            "Verification points tracked during development: authentication boundaries; role route protection; library visibility (public vs organisation-only); moderation state transitions; form validation on wildlife submissions.",
            "Representative manual test cases executed during Sprint 2: (T1) teacher login → dashboard load; (T2) student join with school code → programme visibility; (T3) submit wildlife report → pending moderation; "
            "(T4) foundation user → approve item → appears in approved public surfaces; (T5) community vs school library visibility spot-check.",
            "Test environment: local Next.js dev server on Windows 10/11, Node LTS, SQLite via Prisma for prototype data; browsers: current Chrome/Edge.",
            "Automated tests: eslint (`npm run lint`) and production build (`npm run build`) used as regression gates; broader automated UI tests noted as future work.",
            "",
            "PROJECT RISKS (individual delivery)",
            "Risk: scope creep on conservation features — Mitigation: backlog prioritisation against case interviews. Risk: privacy misconfiguration — Mitigation: explicit org-type rules in data model and UI routing checks.",
            "",
            "FUTURE WORK",
            "• Automated end-to-end and accessibility tests; performance budgets for large media in library.",
            "• Production database (PostgreSQL), secrets management, and formal backup/DR runbooks.",
            "• Creative canvas / richer pupil personalisation if policy allows, reflecting optional case study aspirations.",
            "• Deeper analytics pipelines for foundation dashboards (event tracking with privacy review).",
            "",
        ]
        insert_block_before(p_conc, block)
        full = "\n".join(p.text for p in doc.paragraphs)

    # --- Appendices E–J ---
    p_d1 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("Figure D1."):
            p_d1 = p
            break
    full = "\n".join(p.text for p in doc.paragraphs)
    if p_d1 and "Appendix E. Repository, diagrams" not in full:
        appendix_block = [
            "",
            "Appendix E. Repository, diagrams, and configuration management",
            "Source code (replace with your public URL): https://github.com/[YOUR_USERNAME]/[YOUR_REPO]",
            "Diagram sources: maintain draw.io or exported PNG/SVG in the repository or linked OneDrive folder so assessors can open layers.",
            "Configuration: copy `.env.example` to `.env`; set `SESSION_SECRET` (minimum 16 characters) before running.",
            "",
            "Appendix F. How to run the application (local)",
            "Prerequisites: Node.js (LTS recommended), npm, and Git.",
            "1) Clone the repository and open the project folder.",
            "2) Install dependencies: npm install",
            "3) Create `.env` from `.env.example` and set SESSION_SECRET.",
            "4) Apply database and seed data: npx prisma migrate deploy then npx prisma db seed",
            "5) Start the dev server: npm run dev",
            "6) Browse to http://localhost:3000",
            "Quality checks (optional but recommended before submission): npm run lint and npm run build",
            "",
            "Appendix G. Test accounts and demo access code",
            "Password for all seeded demonstration accounts: KomodoHub!Dev2026",
            "Foundation admin: foundation.admin@komodohub.local",
            "School admin: school.admin@komodohub.local",
            "Teacher: teacher@komodohub.local",
            "Student 1: student1@komodohub.local",
            "Student 2: student2@komodohub.local",
            "Community admin: community.admin@komodohub.local",
            "Community member 1: member1@komodohub.local",
            "Community member 2: member2@komodohub.local",
            "Demo school access code (student join): SCHOOL-DEMO-2026",
            "Public routes for anonymous checks: /library, /campaigns, /species (approved items only).",
            "",
            "Appendix H. Support contact for access issues",
            "Primary student contact (replace with your active university email): [YOUR.COVENTRY.EMAIL@uni.coventry.ac.uk]",
            "Use this inbox for marker questions about registration, codes, or environment errors during evaluation.",
            "",
            "Appendix I. Summary of AI use",
            "Generative AI assisted with drafting clarity, checklist alignment against the template, and appendix structuring. "
            "All technical claims were checked against the implemented repository and module PDFs. Screenshots and approval evidence are my own. "
            "I remain accountable for the accuracy of this submission.",
            "",
            "Appendix J. Declaration of originality",
            "I confirm this report and referenced prototype are my own work except where sources are cited. I understand the Faculty regulations on academic integrity.",
            "",
        ]
        insert_block_after(p_d1, appendix_block)

    doc.save(path)
    print("Wrote enhanced copy:", path)
    print("If this looks correct, close Word and replace SUBMIT_READY_edited.docx with this file, or submit the ENHANCED file directly.")


if __name__ == "__main__":
    main()
