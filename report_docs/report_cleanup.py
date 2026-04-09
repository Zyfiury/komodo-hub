"""Shared cleanup for 5005CMD report DOCX: weekly header dedupe, stray lines, duplicate references."""
from __future__ import annotations

import re

from docx import Document


def _is_sprint_week_heading(text: str) -> bool:
    """True only for real section headings like 'SPRINT 1 – WEEK 1', not backlog lines mentioning Sprint/Week."""
    t = text.strip().upper().replace("–", "-").replace("—", "-")
    return t.startswith("SPRINT") and "WEEK" in t


def _is_week_signature_line(text: str) -> bool:
    t = text.strip()
    if not t.startswith("Omar Zakhama"):
        return False
    return "2026" in t and ("February" in t or "March" in t)


def dedupe_weekly_identity_lines(doc: Document) -> int:
    """Within each SPRINT … WEEK block, keep the first 'Omar Zakhama – role – date' line; blank later identical copies."""
    removed = 0
    seen_in_week: set[str] = set()
    for p in doc.paragraphs:
        raw = p.text.strip()
        if _is_sprint_week_heading(raw):
            seen_in_week = set()
            continue
        if _is_week_signature_line(raw):
            if raw in seen_in_week:
                p.text = ""
                removed += 1
            else:
                seen_in_week.add(raw)
    return removed


def remove_stray_minimal_paragraphs(doc: Document) -> int:
    """Remove paragraphs that are only punctuation (e.g. orphaned '.')."""
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if len(t) <= 2 and t in {".", "..", "-", "—", "–"}:
            p.text = ""
            n += 1
    return n


def dedupe_reference_entries(doc: Document) -> int:
    """Remove exact-duplicate non-empty reference lines between REFERENCES and APPENDICES."""
    in_refs = False
    seen: set[str] = set()
    removed = 0
    for p in doc.paragraphs:
        raw = p.text.strip()
        if raw == "REFERENCES":
            in_refs = True
            seen = set()
            continue
        if in_refs and raw == "APPENDICES":
            break
        if not in_refs or not raw:
            continue
        key = re.sub(r"\s+", " ", raw)
        if key in seen:
            p.text = ""
            removed += 1
        else:
            seen.add(key)
    return removed


def clean_report_document(doc: Document) -> dict[str, int]:
    stats = {
        "deduped_weekly_lines": dedupe_weekly_identity_lines(doc),
        "stray_removed": remove_stray_minimal_paragraphs(doc),
        "deduped_references": dedupe_reference_entries(doc),
    }
    return stats
